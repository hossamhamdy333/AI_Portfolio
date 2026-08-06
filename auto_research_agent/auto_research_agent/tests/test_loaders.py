import os

import docx as docx_lib
import openpyxl
import pytest
from pypdf import PdfWriter

from ingestion.loaders import load_any, load_docx, load_pdf, load_xlsx


@pytest.fixture
def tmp_docx(tmp_path):
    path = tmp_path / "sample.docx"
    d = docx_lib.Document()
    d.add_paragraph("Hello from a Word document.")
    d.add_paragraph("Second paragraph here.")
    d.save(path)
    return str(path)


@pytest.fixture
def tmp_xlsx(tmp_path):
    path = tmp_path / "sample.xlsx"
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Sheet1"
    ws.append(["name", "value"])
    ws.append(["a", 1])
    wb.save(path)
    return str(path)


@pytest.fixture
def tmp_pdf(tmp_path):
    path = tmp_path / "sample.pdf"
    writer = PdfWriter()
    writer.add_blank_page(width=200, height=200)
    with open(path, "wb") as f:
        writer.write(f)
    return str(path)


def test_load_docx_extracts_paragraphs(tmp_docx):
    result = load_docx(tmp_docx)
    assert len(result) == 1
    assert "Hello from a Word document." in result[0]["text"]
    assert result[0]["source"] == "sample.docx"
    assert result[0]["modality"] == "text"


def test_load_xlsx_extracts_rows(tmp_xlsx):
    result = load_xlsx(tmp_xlsx)
    assert len(result) == 1
    assert "Sheet1" in result[0]["text"]
    assert "name, value" in result[0]["text"]
    assert result[0]["modality"] == "table"


def test_load_pdf_blank_page_returns_no_text_chunks(tmp_pdf):
    # A blank page has no extractable text, so the loader should skip it
    # rather than emit an empty chunk.
    result = load_pdf(tmp_pdf)
    assert result == []


def test_load_any_dispatches_by_extension(tmp_docx):
    result = load_any(tmp_docx)
    assert result == load_docx(tmp_docx)


def test_load_any_unknown_extension_raises(tmp_path):
    bad_file = tmp_path / "sample.xyz"
    bad_file.write_text("nope")
    with pytest.raises(ValueError):
        load_any(str(bad_file))
